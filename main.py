from datetime import date
from re import template
from docx import Document
from templates import *
import json
import csv
import sys
import os


MARKER_CONVERSION_DICT = {r"$RECRUITER$": "Recruiter Name",
                          r"$FIRST$": "First Name",
                          r"$LAST$": "Last Name",
                          r"$POSITION$": "Position Title",
                          r"$EMAIL$": "Email"}

def fill_template(parameter_dict: dict):

    # Create empty document object
    document = Document()
    text = ""

    # fill copy with parameters
    text = introductions[parameter_dict["Intro Type"]]
    for key in MARKER_CONVERSION_DICT:
        while key in text:
            text = text.replace(key, parameter_dict[MARKER_CONVERSION_DICT[key]])
    document.add_paragraph().add_run(text)

    
    text = bodies[parameter_dict["Body Type"]]
    for key in MARKER_CONVERSION_DICT:
        while key in text:
            text = text.replace(key, parameter_dict[MARKER_CONVERSION_DICT[key]])
    document.add_paragraph().add_run(text)

    text = concluding[parameter_dict["Conclusion Type"]]
    for key in MARKER_CONVERSION_DICT:
        while key in text:
            text = text.replace(key, parameter_dict[MARKER_CONVERSION_DICT[key]])
    document.add_paragraph().add_run(text)

    # Save Copy
    os.chdir("output")
    name_friendly_ = parameter_dict["Company Name"].replace(" ", "_")
    document.save(f"{name_friendly_}.docx")

def main():

    with open("config.json", "r") as json_file:
        config = json.load(json_file)

    # If "-b" (batch) then enter loop over each row in batches.csv
    if sys.argv[1] == "-b":
    
        with open("batches.csv", "r") as csv_file:
            csv_reader = csv.reader(csv_file)

            for row in csv_reader:

                parameters = {"Company Name": row[0],
                              "Letter Type": row[1],
                              "Recuiter Name": row[2],
                              "Position Title": row[3],
                              "Current Date": f"{date.today().month}/{date.today().day}/{date.today().year}"}

                fill_template(parameters)

    # If "-s" (single) then fill template by hard-coded dictionary
    elif sys.argv[1] == "-s":

        parameters = {"Company Name": "TEST_COMPANY_NAME",
                        "Intro Type": "Test Intro",
                        "Body Type": "Test Body",
                        "Conclusion Type": "Test Conclusion",
                        "Recruiter Name": "RECRUITER_NAME",
                        "Position Title": "POSITION_TITLE",
                        "Current Date": f"{date.today().month}/{date.today().day}/{date.today().year}"}
    
        parameters.update(config)

        fill_template(parameters)

    # Unknown parameter for Cover Letter generation
    else:
        print("\nError: unknown execution param, either \"-s\" or \"-b\"...\n")

if __name__ == "__main__":
    
    main()
